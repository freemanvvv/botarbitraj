"""
Генерация документов: ТЗ, пояснительные записки → DOCX, PDF.
"""
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .config import OUTPUT_DIR


def markdown_to_docx(markdown_text: str, output_path: str | None = None) -> str:
    """
    Конвертирует Markdown в DOCX.
    Базовая реализация без полного парсинга Markdown.
    """
    doc = Document()

    # Заголовок документа
    title = doc.add_heading("Construction AI Copilot", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Дата
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_para.add_run(datetime.now().strftime("%d.%m.%Y")).font.size = Pt(10)

    doc.add_paragraph()  # отступ

    # Парсинг строк (простой — без полного Markdown)
    for line in markdown_text.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue

        # Заголовки
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("> "):
            p = doc.add_paragraph(line[2:])
            p.style = doc.styles["Normal"]
            # курсив
            for run in p.runs:
                run.italic = True
        else:
            doc.add_paragraph(line)

    if output_path is None:
        output_path = os.path.join(OUTPUT_DIR, f"document_{datetime.now():%Y%m%d_%H%M%S}.docx")

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path


def generate_docx(text: str, title: str = "Документ", output_path: str | None = None) -> str:
    """
    Генерирует DOCX из текста (с заголовком).
    """
    full_text = f"# {title}\n\n{text}"
    return markdown_to_docx(full_text, output_path)


def generate_pdf(text: str, title: str = "Документ", output_path: str | None = None) -> str:
    """
    Генерирует PDF из текста через WeasyPrint.
    """
    from weasyprint import HTML

    if output_path is None:
        output_path = os.path.join(
            OUTPUT_DIR, f"pdf_{datetime.now():%Y%m%d_%H%M%S}.pdf"
        )

    # Простая HTML-разметка
    html_body = text.replace("\n", "<br>\n")
    html = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
body {{ font-family: 'DejaVu Sans', sans-serif; font-size: 12pt; margin: 2cm; }}
h1 {{ font-size: 18pt; color: #2c3e50; }}
h2 {{ font-size: 14pt; color: #34495e; }}
h3 {{ font-size: 12pt; color: #7f8c8d; }}
blockquote {{ border-left: 3px solid #ccc; padding-left: 10px; color: #555; }}
</style>
</head>
<body>
<h1>{title}</h1>
<p>{html_body}</p>
</body>
</html>"""

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    HTML(string=html).write_pdf(output_path)
    return output_path
